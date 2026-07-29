"""Generate synthetic, redistributable DOCX regression fixtures for import tests."""

from pathlib import Path
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw

from generate_standard_template import add_real_decimal_numbering, set_numbering


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "app" / "src" / "test" / "resources" / "fixtures"
WORK = ROOT / "build" / "fixture-images"


def make_image(path: Path, color: tuple[int, int, int], label: str):
    image = Image.new("RGB", (96, 64), color)
    draw = ImageDraw.Draw(image)
    draw.rectangle((2, 2, 93, 61), outline=(30, 30, 30), width=2)
    draw.text((40, 24), label, fill=(255, 255, 255))
    image.save(path)


def option(doc: Document, key: str, text: str, pictures=()):
    p = doc.add_paragraph(f"{key}. {text}")
    for picture in pictures:
        p.add_run().add_picture(str(picture), width=Inches(0.45))
    return p


def standard_fixture(images):
    doc = Document()
    num_id = add_real_decimal_numbering(doc)
    doc.add_paragraph("QuizForge synthetic standard regression fixture")

    q = doc.add_paragraph("根据两张示意图选择正确项")
    set_numbering(q, num_id)
    q.add_run().add_picture(str(images[0]), width=Inches(0.45))
    q.add_run().add_picture(str(images[1]), width=Inches(0.45))
    option(doc, "A", "图形甲", pictures=images[2:4])
    option(doc, "B", "图形乙")
    doc.add_paragraph("答案：A")
    doc.add_paragraph("解析：题干和选项均包含多张图片。")
    doc.add_paragraph("知识点：图片题")

    doc.add_paragraph("2、下列属于传输层协议的是？")
    option(doc, "A", "TCP")
    option(doc, "B", "UDP")
    option(doc, "C", "IP")
    doc.add_paragraph("答案：A、B")

    doc.add_paragraph("（3）进程和程序完全相同。")
    doc.add_paragraph("答案：错")

    doc.add_paragraph("第4题 这道题故意缺少答案")
    option(doc, "A", "甲")
    option(doc, "B", "乙")
    doc.save(OUT / "standard-regression.docx")


def smart_fixture():
    doc = Document()
    doc.add_paragraph("第一章 混合格式练习")
    doc.add_paragraph(
        "1. JVM 运行哪种字节码？ A. Java bytecode B. Native code 答案：A "
        "2、Kotlin 默认运行在哪个平台？ A、JVM B、BIOS 答案：A"
    )
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "3. 表格中的题目"
    table.cell(0, 1).text = "A. 甲\nB. 乙"
    table.cell(1, 0).text = "答案"
    table.cell(1, 1).text = "B"
    doc.add_paragraph("第二节 标题夹在题目中")
    doc.add_paragraph("4. 集中答案题 A. 甲 B. 乙")
    doc.add_paragraph("5. 另一道集中答案题 A. 对 B. 错")
    doc.add_paragraph("答案汇总：4.A 5.B")
    doc.add_paragraph("超长说明：" + ("跨块上下文内容。" * 900))
    doc.save(OUT / "smart-regression.docx")


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    WORK.mkdir(parents=True, exist_ok=True)
    specs = [
        ((39, 112, 180), "1"),
        ((46, 139, 87), "2"),
        ((164, 92, 38), "A"),
        ((117, 71, 158), "B"),
    ]
    images = []
    for index, (color, label) in enumerate(specs):
        path = WORK / f"fixture-{index}.png"
        make_image(path, color, label)
        images.append(path)
    standard_fixture(images)
    smart_fixture()
    print(OUT)


if __name__ == "__main__":
    main()
