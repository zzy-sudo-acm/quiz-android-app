"""Generate the bundled QuizForge standard-format Word template."""

from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "app" / "src" / "main" / "assets" / "quizforge-standard-template.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x55, 0x55, 0x55)


def set_font(run, size=11, color=None, bold=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Title": (22, DARK_BLUE, 0, 8),
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_real_decimal_numbering(doc):
    numbering = doc.part.numbering_part.element
    existing_abstract = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(existing_abstract or [0]) + 1
    existing_nums = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    num_id = max(existing_nums or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "decimal")
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "%1.")
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.extend([tabs, indent, spacing])
    level.extend([start, fmt, text, justification, ppr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def set_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def add_option(doc, key, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(0)
    label = paragraph.add_run(f"{key}. ")
    set_font(label, bold=True, color=DARK_BLUE)
    set_font(paragraph.add_run(text))


def add_field(doc, label, value, color=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.375)
    run = paragraph.add_run(f"{label}：")
    set_font(run, bold=True, color=color or DARK_BLUE)
    set_font(paragraph.add_run(value))


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    set_font(paragraph.add_run(text))


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    num_id = add_real_decimal_numbering(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(title.add_run("QuizForge 标准题库模板"), size=22, color=DARK_BLUE, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    set_font(subtitle.add_run("完全离线导入 · 无需 API · 支持单选、多选、判断、解析、知识点和图片"), size=10.5, color=MUTED)

    doc.add_heading("使用说明", level=1)
    add_bullet(doc, "每道题必须有明确题号；本模板使用 Word 自动编号，复制题目段落即可继续编号。")
    add_bullet(doc, "每个选项建议单独一段，支持 A.、A、A：、A:、A) 和 A）等标记。")
    add_bullet(doc, "答案必须放在对应题目后；解析和知识点可以省略。")
    add_bullet(doc, "图片请插入题干段落或对应选项段落；一题和一个选项均可插入多张图片。")
    add_bullet(doc, "不要删除“答案：”字段；格式错误不会自动调用 AI，也不会静默跳过。")

    doc.add_heading("标准示例", level=1)
    question = doc.add_paragraph()
    set_numbering(question, num_id)
    set_font(question.add_run("下列关于操作系统的说法正确的是？"), bold=True)
    add_option(doc, "A", "操作系统只负责运行应用程序")
    add_option(doc, "B", "操作系统负责管理计算机软硬件资源")
    add_option(doc, "C", "操作系统不管理内存")
    add_option(doc, "D", "操作系统只存在于电脑中")
    add_field(doc, "答案", "B")
    add_field(doc, "解析", "操作系统负责管理和调度计算机软硬件资源。")
    add_field(doc, "知识点", "操作系统基础")

    question = doc.add_paragraph()
    set_numbering(question, num_id)
    set_font(question.add_run("下列属于传输层协议的是？"), bold=True)
    add_option(doc, "A", "TCP")
    add_option(doc, "B", "UDP")
    add_option(doc, "C", "IP")
    add_option(doc, "D", "ARP")
    add_field(doc, "答案", "AB")
    add_field(doc, "解析", "TCP 和 UDP 属于传输层协议。")
    add_field(doc, "知识点", "计算机网络")

    question = doc.add_paragraph()
    set_numbering(question, num_id)
    set_font(question.add_run("进程和程序是完全相同的概念。"), bold=True)
    add_field(doc, "答案", "错")
    add_field(doc, "解析", "程序是静态代码，进程是程序的一次运行过程。")
    add_field(doc, "知识点", "进程管理")

    doc.add_heading("继续编辑", level=1)
    paragraph = doc.add_paragraph()
    set_numbering(paragraph, num_id)
    placeholder = paragraph.add_run("在这里输入下一道题的题干（保留自动编号）")
    set_font(placeholder, bold=True, color=MUTED)
    add_option(doc, "A", "选项 A")
    add_option(doc, "B", "选项 B")
    add_field(doc, "答案", "A")
    add_field(doc, "解析", "可选")
    add_field(doc, "知识点", "可选")

    core = doc.core_properties
    core.title = "QuizForge 标准题库模板"
    core.subject = "QuizForge 完全离线标准格式导入模板"
    core.author = "QuizForge"
    core.keywords = "QuizForge, 题库, DOCX, 标准格式"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
